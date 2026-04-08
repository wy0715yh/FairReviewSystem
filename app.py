import io
import os
import re
import time
from html import unescape
from typing import List

from flask import Flask, jsonify, render_template, request, send_file
from docx import Document

from services.audit_engine import run_batch_audit, run_single_audit
from services.config import ADMIN_PASSWORD, MAX_AUDIT_TARGET_CHARS, MAX_BATCH_ITEMS
from services.llm_client import call_llm, strip_think_blocks
from services.parsers import concat_text_input, read_uploaded_file
from services.rag import search_knowledge
from services.reporting import render_batch_report, render_single_report
from services.storage import (
    add_history,
    clear_history,
    delete_rule,
    list_rules,
    load_custom_rules,
    load_history,
    load_settings,
    remove_history,
    remove_custom_rule,
    save_rule,
    save_settings,
    upsert_custom_rule,
)

app = Flask(__name__)


def _html_to_plain_text(report_html: str) -> str:
    html_text = report_html or ""
    html_text = re.sub(r"<br\s*/?>", "\n", html_text, flags=re.IGNORECASE)
    html_text = re.sub(r"</(p|li|h1|h2|h3|h4|h5|h6|div|tr)>", "\n", html_text, flags=re.IGNORECASE)
    html_text = re.sub(r"<[^>]+>", "", html_text)
    html_text = unescape(html_text)
    html_text = re.sub(r"\n{3,}", "\n\n", html_text)
    return html_text.strip()


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/history')
def history_page():
    return render_template('history.html')


@app.route('/api/verify_admin', methods=['POST'])
def verify_admin():
    data = request.json or {}
    if data.get('password') == ADMIN_PASSWORD:
        return jsonify({'status': 'success'})
    return jsonify({'status': 'fail'}), 403


@app.route('/api/upload_rules', methods=['POST'])
def upload_rules():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'message': '无文件上传'}), 400

    count = 0
    for f in files:
        content = read_uploaded_file(f)
        if not content.strip():
            continue
        stem = os.path.splitext(f.filename)[0].strip() or f'规则{int(time.time())}'
        save_rule(stem, content)
        count += 1
    return jsonify({'message': f'成功入库 {count} 个文件'})


@app.route('/api/list_rules', methods=['GET'])
def api_list_rules():
    return jsonify({'files': list_rules()})


@app.route('/api/delete_rule', methods=['POST'])
def api_delete_rule():
    data = request.json or {}
    filename = str(data.get('filename', '')).strip()
    if not filename:
        return jsonify({'message': '文件名不能为空'}), 400
    if delete_rule(filename):
        return jsonify({'message': f'成功删除 {filename}'})
    return jsonify({'message': '文件不存在'}), 404


@app.route('/api/custom_rules', methods=['GET'])
def api_custom_rules_get():
    return jsonify({'rules': load_custom_rules()})


@app.route('/api/custom_rules', methods=['POST', 'PUT'])
def api_custom_rules_upsert():
    data = request.json or {}
    row = upsert_custom_rule(data)
    return jsonify({'rule': row})


@app.route('/api/custom_rules/<rule_id>', methods=['DELETE'])
def api_custom_rules_delete(rule_id):
    if remove_custom_rule(rule_id):
        return jsonify({'message': '删除成功'})
    return jsonify({'message': '规则不存在'}), 404


@app.route('/api/settings', methods=['GET'])
def api_settings_get():
    return jsonify({'settings': load_settings()})


@app.route('/api/settings', methods=['POST'])
def api_settings_save():
    data = request.json or {}
    saved = save_settings(data)
    return jsonify({'settings': saved})


@app.route('/api/knowledge_search', methods=['GET'])
def api_knowledge_search():
    q = (request.args.get('q') or '').strip()
    if not q:
        return jsonify({'items': []})
    items = search_knowledge(q, top_k=20, max_chunk_chars=420)
    return jsonify({'items': items})


@app.route('/api/audit', methods=['POST'])
def api_audit():
    text = request.form.get('text', '')
    file = request.files.get('file')
    target = concat_text_input(text, [file] if file else [])

    if not target:
        return jsonify({'error': '内容为空'}), 400

    target = target[:MAX_AUDIT_TARGET_CHARS]
    result = run_single_audit(target)
    html = render_single_report(result['structured'], elapsed_seconds=result['elapsed'])

    add_history(
        title=result['structured'].get('policy_name', target[:60]),
        content=target,
        report_html=html,
        structured=result['structured'],
    )

    payload = {
        'result': html,
        'structured': result['structured'],
        'fallback': result['fallback'],
    }
    if result.get('error'):
        payload['error'] = result['error']
    return jsonify(payload)


@app.route('/api/audit_batch', methods=['POST'])
def api_audit_batch():
    texts: List[str] = []

    json_data = request.json if request.is_json else None
    if isinstance(json_data, dict):
        for x in json_data.get('texts', []):
            s = str(x).strip()
            if s:
                texts.append(s)

    form_texts = request.form.get('texts', '')
    if form_texts and form_texts.strip():
        chunks = [x.strip() for x in form_texts.split('\n\n') if x.strip()]
        texts.extend(chunks)

    files = request.files.getlist('files')
    for f in files:
        payload = read_uploaded_file(f).strip()
        if payload:
            texts.append(payload)

    if not texts:
        return jsonify({'error': '请提供批量文本或文件'}), 400
    if len(texts) > MAX_BATCH_ITEMS:
        return jsonify({'error': f'单次批量最多 {MAX_BATCH_ITEMS} 份'}), 400

    texts = [x[:MAX_AUDIT_TARGET_CHARS] for x in texts]
    summary = run_batch_audit(texts)
    html = render_batch_report(summary)

    add_history(
        title=f'批量审查 {summary.get("total", 0)} 份',
        content='\n\n'.join(texts[:3])[:1200],
        report_html=html,
        structured=summary,
    )

    return jsonify({'result': html, 'summary': summary})


@app.route('/api/history', methods=['GET'])
def api_history_list():
    history = load_history()
    simple = [{'id': h['id'], 'title': h['title'], 'time': h['time']} for h in history]
    return jsonify({'history': simple})


@app.route('/api/history/<int:record_id>', methods=['GET'])
def api_history_detail(record_id: int):
    history = load_history()
    for h in history:
        if int(h.get('id', 0)) == int(record_id):
            return jsonify({'record': h})
    return jsonify({'error': '记录不存在'}), 404


@app.route('/api/history/<int:record_id>', methods=['DELETE'])
def api_history_delete(record_id: int):
    if remove_history(record_id):
        return jsonify({'message': '删除成功'})
    return jsonify({'error': '记录不存在'}), 404


@app.route('/api/history', methods=['DELETE'])
def api_history_clear():
    count = clear_history()
    return jsonify({'message': f'已清空 {count} 条记录', 'count': count})


@app.route('/api/chat_ip', methods=['POST'])
def chat_ip():
    data = request.json or {}
    query = str(data.get('query', '')).strip()
    if not query:
        return jsonify({'error': 'query不能为空'}), 400

    prompt = query + "\n\n【输出格式要求】请使用清晰标题和分段，每个要点单独成段。"
    try:
        text = call_llm(prompt, max_total_seconds=25, max_tokens=900)
        clean = strip_think_blocks(text).replace('```html', '').replace('```', '')
        clean = clean.replace('\n', '<br>　　')
        return jsonify({'result': clean})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/export_report', methods=['POST'])
def export_report():
    data = request.json or {}
    fmt = str(data.get('format', '')).strip().lower()
    report_html = str(data.get('html', '')).strip()
    title = str(data.get('title', '审查报告')).strip() or '审查报告'
    if fmt not in ['docx', 'pdf']:
        return jsonify({'error': '仅支持 docx 或 pdf'}), 400
    if not report_html:
        return jsonify({'error': '报告内容为空'}), 400

    plain = _html_to_plain_text(report_html)
    lines = [x.strip() for x in plain.splitlines() if x.strip()]
    if not lines:
        lines = ['（空白报告）']

    if fmt == 'docx':
        doc = Document()
        doc.add_heading(title, level=1)
        for ln in lines:
            doc.add_paragraph(ln)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=f'{title}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except Exception:
        return jsonify({'error': '缺少 reportlab 依赖，请先安装'}), 500

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    c.setFont('STSong-Light', 11)
    width, height = A4
    left = 42
    top = height - 42
    line_h = 16

    def draw_line(y, txt):
        c.drawString(left, y, txt)

    y = top
    draw_line(y, title)
    y -= line_h * 1.5

    max_chars = 44
    for ln in lines:
        chunks = [ln[i:i + max_chars] for i in range(0, len(ln), max_chars)] or [""]
        for chunk in chunks:
            if y < 48:
                c.showPage()
                c.setFont('STSong-Light', 11)
                y = top
            draw_line(y, chunk)
            y -= line_h

    c.save()
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f'{title}.pdf',
        mimetype='application/pdf',
    )


if __name__ == '__main__':
    app.run(debug=False, port=5000, use_reloader=False)
